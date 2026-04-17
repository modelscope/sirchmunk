# Copyright (c) ModelScope Contributors. All rights reserved.
"""
TOC (Table of Contents) extractor — multi-layer fallback strategy.

Extracts hierarchical table-of-contents structures from various document
formats (PDF, Markdown, DOCX, HTML) using a layered approach:

  Layer 1 — pypdf native outline (highest confidence, zero cost)
  Layer 2 — pdfminer.six detailed parsing (fallback for pypdf)
  Layer 3 — Text heading pattern detection (for documents without bookmarks)
  Layer 4 — LLM-assisted inference (optional, last resort)

The extracted TOCEntry list is consumed by the tree indexer to accelerate
tree construction.
"""

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, ClassVar, List, Optional

logger = logging.getLogger(__name__)

# Known heading-style prefixes across locales (English, Chinese, etc.)
_HEADING_STYLE_PREFIXES = ("Heading", "heading", "\u6807\u9898")  # "标题" = Chinese


# ---------------------------------------------------------------------------
#  Data models
# ---------------------------------------------------------------------------


@dataclass
class TOCEntry:
    """Single entry in an extracted table of contents.

    Attributes:
        title: Section title text.
        level: Heading depth (1 = top-level section, 2 = subsection, …).
        char_start: Character offset in the extracted full text.
        char_end: End character offset (exclusive), or None if unresolved.
        page_start: 1-indexed page number, or None if unknown.
        page_end: End page number (inclusive), or None.
        children: Nested sub-entries forming a tree.
        source: Which extraction layer produced this entry
                ("pypdf", "pdfminer", "heading", "markdown", "docx",
                 "html", "llm").
    """

    title: str
    level: int  # 1=section, 2=subsection, …
    char_start: int = 0
    char_end: Optional[int] = None
    page_start: Optional[int] = None
    page_end: Optional[int] = None
    children: List["TOCEntry"] = field(default_factory=list)
    source: str = ""


@dataclass
class TocResult:
    """Complete TOC extraction result with quality metadata.

    Attributes:
        entries: Ordered list of TOCEntry objects.
        source: Primary extraction method that produced the result.
        confidence: Estimated quality score (0.0–1.0).
        page_count: Total pages in the source document, if known.
    """

    entries: List[TOCEntry] = field(default_factory=list)
    source: str = ""
    confidence: float = 0.0
    page_count: Optional[int] = None


# ---------------------------------------------------------------------------
#  Layer 1: pypdf native outline
# ---------------------------------------------------------------------------


class PypdfOutlineExtractor:
    """Layer 1: Extract TOC from PDF native outline/bookmarks using pypdf.

    Highest confidence (0.9) — relies on the PDF producer embedding
    explicit bookmarks.  Zero external cost.
    """

    @staticmethod
    def extract(file_path: str | Path) -> TocResult:
        """Extract TOC from PDF outline.

        Args:
            file_path: Path to the PDF file.

        Returns:
            TocResult with entries and page_count populated,
            or an empty TocResult on failure.
        """
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            outline = reader.outline
            page_count = len(reader.pages)

            if not outline:
                return TocResult(source="pypdf", page_count=page_count)

            entries: List[TOCEntry] = []
            PypdfOutlineExtractor._parse_outline(
                reader, outline, entries, level=1,
            )

            if not entries:
                return TocResult(source="pypdf", page_count=page_count)

            return TocResult(
                entries=entries,
                source="pypdf",
                confidence=0.9,
                page_count=page_count,
            )
        except Exception as exc:
            logger.debug("pypdf outline extraction failed: %s", exc)
            return TocResult(source="pypdf")

    @staticmethod
    def _parse_outline(
        reader: Any,
        outline_items: list,
        entries: List[TOCEntry],
        level: int,
    ) -> None:
        """Recursively parse pypdf outline items into TOCEntry list."""
        for item in outline_items:
            if isinstance(item, list):
                # Nested list → sub-bookmarks; attach to last entry
                if entries:
                    sub: List[TOCEntry] = []
                    PypdfOutlineExtractor._parse_outline(
                        reader, item, sub, level=level + 1,
                    )
                    entries[-1].children.extend(sub)
                else:
                    PypdfOutlineExtractor._parse_outline(
                        reader, item, entries, level=level,
                    )
            else:
                try:
                    title = item.title if hasattr(item, "title") else str(item)
                    page_num: Optional[int] = None
                    try:
                        # get_destination_page_number returns 0-indexed
                        raw = reader.get_destination_page_number(item)
                        if raw is not None:
                            page_num = raw + 1  # convert to 1-indexed
                    except Exception:
                        pass
                    entries.append(TOCEntry(
                        title=title.strip(),
                        level=level,
                        char_start=0,
                        page_start=page_num,
                        source="pypdf",
                    ))
                except Exception:
                    continue


# ---------------------------------------------------------------------------
#  Layer 2: pdfminer.six detailed parsing
# ---------------------------------------------------------------------------


class PdfminerOutlineExtractor:
    """Layer 2: Extract TOC using pdfminer.six for more detailed parsing.

    Falls back here when pypdf yields insufficient entries.
    Confidence 0.85 — pdfminer exposes more detail but requires
    manual page-number resolution.
    """

    @staticmethod
    def extract(file_path: str | Path) -> TocResult:
        """Extract TOC using pdfminer's outline parser.

        Args:
            file_path: Path to the PDF file.

        Returns:
            TocResult with entries populated, or empty on failure.
        """
        try:
            from pdfminer.pdfdocument import PDFDocument, PDFNoOutlines
            from pdfminer.pdfpage import PDFPage
            from pdfminer.pdfparser import PDFParser
            from pdfminer.psparser import LIT

            fp = open(str(file_path), "rb")
            try:
                parser = PDFParser(fp)
                document = PDFDocument(parser)

                # Build page-object-id → 1-indexed page number mapping
                pages = list(PDFPage.create_pages(document))
                page_count = len(pages)
                objid_to_pagenum = {
                    page.pageid: idx + 1
                    for idx, page in enumerate(pages)
                }

                entries: List[TOCEntry] = []
                try:
                    for level, title, dest, action, _se in document.get_outlines():
                        page_num = PdfminerOutlineExtractor._resolve_page(
                            dest, action, objid_to_pagenum, document,
                        )
                        entries.append(TOCEntry(
                            title=str(title).strip() if title else "",
                            level=level,
                            char_start=0,
                            page_start=page_num,
                            source="pdfminer",
                        ))
                except PDFNoOutlines:
                    pass

                if not entries:
                    return TocResult(source="pdfminer", page_count=page_count)

                return TocResult(
                    entries=entries,
                    source="pdfminer",
                    confidence=0.85,
                    page_count=page_count,
                )
            finally:
                fp.close()
        except Exception as exc:
            logger.debug("pdfminer outline extraction failed: %s", exc)
            return TocResult(source="pdfminer")

    @staticmethod
    def _resolve_page(
        dest: Any,
        action: Any,
        objid_to_pagenum: dict,
        document: Any,
    ) -> Optional[int]:
        """Resolve a pdfminer destination/action to a 1-indexed page number."""
        try:
            from pdfminer.pdfparser import PDFStream
            from pdfminer.pdftypes import resolve1

            # Try dest first
            target = dest
            if target is None and action is not None:
                # GoTo action: action dict may have a 'D' key
                if isinstance(action, dict):
                    target = action.get("D")

            if target is None:
                return None

            # Resolve indirect objects
            target = resolve1(target)

            if isinstance(target, list) and len(target) > 0:
                page_ref = resolve1(target[0])
                if hasattr(page_ref, "objid"):
                    return objid_to_pagenum.get(page_ref.objid)
            elif hasattr(target, "objid"):
                return objid_to_pagenum.get(target.objid)
        except Exception:
            pass
        return None


# ---------------------------------------------------------------------------
#  Layer 3: Text heading pattern detection
# ---------------------------------------------------------------------------


class HeadingTocExtractor:
    """Layer 3: Infer TOC from document text structure (heading patterns).

    Handles Markdown headings, numbered sections, and common structural
    keywords.  Confidence 0.6 — heuristic-based, lower precision.
    """

    # Regex for Markdown ATX headings: # Title, ## Subtitle, …
    _MD_HEADING_RE: ClassVar[re.Pattern] = re.compile(
        r"^(#{1,6})\s+(.+)$", re.MULTILINE,
    )

    # Regex for numbered section patterns: "1.", "1.1", "1.1.1", …
    _NUMBERED_RE: ClassVar[re.Pattern] = re.compile(
        r"^(\d+(?:\.\d+)*)[.\s]+(.+)$", re.MULTILINE,
    )

    # Common structural keywords (case-insensitive)
    _STRUCTURAL_KEYWORDS: ClassVar[tuple] = (
        "ITEM", "PART", "CHAPTER", "SECTION", "ARTICLE",
        "APPENDIX", "EXHIBIT", "SCHEDULE", "ANNEX",
    )

    # Max characters for a candidate heading line
    _MAX_HEADING_LINE_LEN: ClassVar[int] = 120

    @staticmethod
    def extract(content: str, mime_type: str = "") -> TocResult:
        """Infer TOC from text content by detecting heading patterns.

        Tries strategies in order:
          1. Markdown ATX headings (``#`` syntax)
          2. Numbered section patterns (``1.``, ``1.1``, …)
          3. Structural keyword detection (ITEM, PART, CHAPTER, …)

        Args:
            content: Full extracted text of the document.
            mime_type: Optional MIME type hint (unused currently).

        Returns:
            TocResult with char_position-based entries.
        """
        if not content or len(content.strip()) < 50:
            return TocResult(source="heading")

        # Strategy 1: Markdown headings
        entries = HeadingTocExtractor._extract_markdown_headings(content)
        if entries:
            return TocResult(
                entries=entries,
                source="heading",
                confidence=0.7,
            )

        # Strategy 2: Numbered sections
        entries = HeadingTocExtractor._extract_numbered_sections(content)
        if entries:
            return TocResult(
                entries=entries,
                source="heading",
                confidence=0.6,
            )

        # Strategy 3: Structural keywords + heuristic
        entries = HeadingTocExtractor._extract_structural_headings(content)
        if entries:
            return TocResult(
                entries=entries,
                source="heading",
                confidence=0.5,
            )

        return TocResult(source="heading")

    @staticmethod
    def _extract_markdown_headings(content: str) -> List[TOCEntry]:
        """Extract headings from Markdown ATX syntax (# / ## / ###)."""
        matches = list(HeadingTocExtractor._MD_HEADING_RE.finditer(content))
        if not matches:
            return []

        entries: List[TOCEntry] = []
        for m in matches:
            hashes, title = m.group(1), m.group(2).strip()
            if title:
                entries.append(TOCEntry(
                    title=title,
                    level=len(hashes),
                    char_start=m.start(),
                    source="heading",
                ))
        return entries

    @staticmethod
    def _extract_numbered_sections(content: str) -> List[TOCEntry]:
        """Extract numbered section headings (1., 1.1, 1.1.1, …)."""
        matches = list(HeadingTocExtractor._NUMBERED_RE.finditer(content))
        if not matches:
            return []

        entries: List[TOCEntry] = []
        for m in matches:
            number_part = m.group(1)
            title_part = m.group(2).strip()
            # Line length check — skip long lines (likely not headings)
            line_len = m.end() - m.start()
            if line_len > HeadingTocExtractor._MAX_HEADING_LINE_LEN:
                continue
            if not title_part:
                continue
            level = number_part.count(".") + 1
            entries.append(TOCEntry(
                title=f"{number_part} {title_part}",
                level=level,
                char_start=m.start(),
                source="heading",
            ))
        return entries

    @staticmethod
    def _extract_structural_headings(content: str) -> List[TOCEntry]:
        """Detect common structural keywords as section boundaries."""
        # Build pattern: ITEM 1, PART I, CHAPTER 1, etc.
        kw_pattern = "|".join(HeadingTocExtractor._STRUCTURAL_KEYWORDS)
        pattern = re.compile(
            rf"^({kw_pattern})\s+(\w+[\w .:\-]*)$",
            re.MULTILINE | re.IGNORECASE,
        )
        matches = list(pattern.finditer(content))
        if not matches:
            return []

        entries: List[TOCEntry] = []
        for m in matches:
            keyword = m.group(1).upper()
            rest = m.group(2).strip()
            title = f"{keyword} {rest}"
            # Determine level based on keyword
            if keyword in ("PART", "CHAPTER"):
                level = 1
            elif keyword in ("ITEM", "SECTION", "ARTICLE"):
                level = 2
            else:
                level = 3
            entries.append(TOCEntry(
                title=title,
                level=level,
                char_start=m.start(),
                source="heading",
            ))
        return entries


# ---------------------------------------------------------------------------
#  Layer 4: LLM-assisted inference (optional)
# ---------------------------------------------------------------------------


class LlmTocExtractor:
    """Layer 4: Use LLM to infer TOC from document content.

    This is the last-resort fallback.  Requires an ``llm_caller`` that
    supports ``await llm_caller.achat(messages)``.  If no caller is
    provided, returns an empty result immediately.

    Confidence 0.7 — LLM may hallucinate structure.
    """

    # Maximum characters sent to the LLM to stay within token limits
    _MAX_CONTENT_CHARS: ClassVar[int] = 8_000

    _PROMPT_TEMPLATE: ClassVar[str] = (
        "Analyze the following document excerpt and extract its "
        "hierarchical table of contents (TOC) structure.\n\n"
        "Return a JSON array where each element has:\n"
        '  - "title": section title text\n'
        '  - "level": integer heading depth (1=top, 2=sub, 3=subsub)\n\n'
        "Only include actual section/chapter headings, not every paragraph.\n"
        "Return ONLY the JSON array, no other text.\n\n"
        "Document excerpt:\n---\n{content}\n---"
    )

    @staticmethod
    async def extract(
        content: str,
        llm_caller: Any | None = None,
    ) -> TocResult:
        """Infer TOC using LLM analysis.

        Args:
            content: Full extracted text of the document.
            llm_caller: An object with ``achat(messages)`` method.
                        If None, returns an empty result.

        Returns:
            TocResult with LLM-inferred entries.
        """
        if llm_caller is None:
            return TocResult(source="llm")

        if not content or len(content.strip()) < 100:
            return TocResult(source="llm")

        try:
            # Truncate content to fit token budget
            truncated = content[:LlmTocExtractor._MAX_CONTENT_CHARS]
            prompt = LlmTocExtractor._PROMPT_TEMPLATE.format(content=truncated)

            resp = await llm_caller.achat([{"role": "user", "content": prompt}])
            raw = resp.content.strip()

            entries = LlmTocExtractor._parse_response(raw, content)
            if not entries:
                return TocResult(source="llm")

            return TocResult(
                entries=entries,
                source="llm",
                confidence=0.7,
            )
        except Exception as exc:
            logger.debug("LLM TOC extraction failed: %s", exc)
            return TocResult(source="llm")

    @staticmethod
    def _parse_response(raw: str, content: str) -> List[TOCEntry]:
        """Parse LLM JSON response into TOCEntry list with char_positions."""
        # Strip markdown code fences if present
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            lines = cleaned.split("\n")
            # Remove first and last fence lines
            lines = [l for l in lines if not l.strip().startswith("```")]
            cleaned = "\n".join(lines)

        try:
            items = json.loads(cleaned)
        except (json.JSONDecodeError, TypeError):
            return []

        if not isinstance(items, list):
            return []

        content_lower = content.lower()
        search_from = 0
        entries: List[TOCEntry] = []

        for item in items:
            if not isinstance(item, dict):
                continue
            title = str(item.get("title", "")).strip()
            level = int(item.get("level", 1))
            if not title:
                continue

            # Try to locate title in content for char_position
            pos = content_lower.find(title.lower(), search_from)
            if pos >= 0:
                char_start = pos
                search_from = pos + len(title)
            else:
                # Fallback: try from beginning
                pos = content_lower.find(title.lower())
                char_start = pos if pos >= 0 else search_from

            entries.append(TOCEntry(
                title=title,
                level=max(1, min(level, 6)),
                char_start=char_start,
                source="llm",
            ))

        return entries


# ---------------------------------------------------------------------------
#  Format-specific extractors (non-PDF)
# ---------------------------------------------------------------------------


class DocxTocExtractor:
    """Extract TOC from DOCX heading styles using python-docx."""

    @staticmethod
    def extract(file_path: str | Path) -> TocResult:
        """Extract TOC from DOCX heading styles.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            TocResult with entries from heading styles.
        """
        try:
            import docx

            doc = docx.Document(str(file_path))
            entries: List[TOCEntry] = []
            for para in doc.paragraphs:
                style_name = para.style.name or ""
                matched_prefix = ""
                for prefix in _HEADING_STYLE_PREFIXES:
                    if style_name.startswith(prefix):
                        matched_prefix = prefix
                        break
                if not matched_prefix:
                    continue
                level_str = style_name[len(matched_prefix):].strip()
                try:
                    level = int(level_str) if level_str else 1
                except ValueError:
                    level = 1
                title = para.text.strip()
                if title:
                    entries.append(TOCEntry(
                        title=title,
                        level=level,
                        char_start=0,
                        source="docx",
                    ))

            if not entries:
                return TocResult(source="docx")
            return TocResult(entries=entries, source="docx", confidence=0.85)
        except Exception as exc:
            logger.debug("DOCX TOC extraction failed: %s", exc)
            return TocResult(source="docx")


class HtmlTocExtractor:
    """Extract TOC from HTML heading tags (<h1>–<h6>)."""

    _HTML_HEADING_RE: ClassVar[re.Pattern] = re.compile(
        r"<h([1-6])[^>]*>(.*?)</h\1>",
        re.IGNORECASE | re.DOTALL,
    )

    @staticmethod
    def extract(content: str) -> TocResult:
        """Extract TOC from HTML heading tags.

        Args:
            content: HTML text content.

        Returns:
            TocResult with entries from <h1>–<h6> tags.
        """
        try:
            matches = HtmlTocExtractor._HTML_HEADING_RE.findall(content)
            if not matches:
                return TocResult(source="html")

            entries: List[TOCEntry] = []
            for level_str, raw_title in matches:
                title = re.sub(r"<[^>]+>", "", raw_title).strip()
                if title:
                    entries.append(TOCEntry(
                        title=title,
                        level=int(level_str),
                        char_start=0,
                        source="html",
                    ))

            if not entries:
                return TocResult(source="html")
            return TocResult(entries=entries, source="html", confidence=0.8)
        except Exception as exc:
            logger.debug("HTML TOC extraction failed: %s", exc)
            return TocResult(source="html")


# ---------------------------------------------------------------------------
#  Orchestrator: multi-layer fallback
# ---------------------------------------------------------------------------


class TOCExtractor:
    """Orchestrates multi-layer TOC extraction with fallback strategy.

    All methods are static/classmethod — no instance state required.
    The main ``extract()`` entry point dispatches by file extension and
    applies the layered fallback for PDF files.

    Layer priority for PDFs:
      1. pypdf native outline (confidence 0.9)
      2. pdfminer.six detailed parsing (confidence 0.85)
      3. Text heading detection (confidence 0.5–0.7)
      4. LLM-assisted inference (confidence 0.7, optional)

    Design constraints:
      - Layers 1–3 are pure-local, zero LLM calls
      - Layer 4 is optional (requires llm_caller)
      - Each layer is independently try-excepted; failure never blocks
        subsequent layers
    """

    # Minimum entries to consider a TOC extraction successful
    _MIN_ENTRIES_THRESHOLD: ClassVar[int] = 3

    @classmethod
    async def extract(
        cls,
        file_path: str,
        content: str,
        *,
        llm_caller: Any | None = None,
    ) -> Optional[List[TOCEntry]]:
        """Extract TOC using layered fallback strategy.

        Tries extraction methods in order of reliability.  Falls back to
        the next layer when the current layer yields fewer than
        ``_MIN_ENTRIES_THRESHOLD`` entries.

        Args:
            file_path: Absolute path to the source file.
            content: Extracted text content of the file.
            llm_caller: Optional LLM caller for Layer 4.

        Returns:
            List of TOCEntry with resolved char positions, or None if
            no layer produced enough entries.
        """
        ext = Path(file_path).suffix.lower()

        result: Optional[TocResult] = None

        if ext == ".pdf":
            result = await cls._extract_pdf_layered(
                file_path, content, llm_caller,
            )
        elif ext in (".md", ".markdown"):
            heading_result = HeadingTocExtractor.extract(content)
            if cls._is_sufficient(heading_result):
                result = heading_result
        elif ext in (".docx",):
            result = DocxTocExtractor.extract(file_path)
        elif ext in (".html", ".htm"):
            result = HtmlTocExtractor.extract(content)
        else:
            return None

        if result is None or not cls._is_sufficient(result):
            return None

        entries = result.entries
        total = cls._count_entries(entries)
        if total < cls._MIN_ENTRIES_THRESHOLD:
            return None

        # Resolve character positions in the extracted text
        entries = cls._resolve_char_positions(entries, content)
        return entries

    @classmethod
    async def _extract_pdf_layered(
        cls,
        file_path: str,
        content: str,
        llm_caller: Any | None,
    ) -> Optional[TocResult]:
        """Apply layered extraction for PDF files.

        Args:
            file_path: Path to the PDF file.
            content: Extracted text content.
            llm_caller: Optional LLM caller for Layer 4.

        Returns:
            Best TocResult from the layer cascade, or None.
        """
        # Layer 1: pypdf
        result = PypdfOutlineExtractor.extract(file_path)
        if cls._is_sufficient(result):
            logger.info(
                "TOC Layer 1 (pypdf): %d entries for %s",
                len(result.entries), Path(file_path).name,
            )
            return result

        # Layer 2: pdfminer.six
        result = PdfminerOutlineExtractor.extract(file_path)
        if cls._is_sufficient(result):
            logger.info(
                "TOC Layer 2 (pdfminer): %d entries for %s",
                len(result.entries), Path(file_path).name,
            )
            return result

        # Layer 3: heading detection from content
        if content:
            result = HeadingTocExtractor.extract(content)
            if cls._is_sufficient(result):
                logger.info(
                    "TOC Layer 3 (heading): %d entries for %s",
                    len(result.entries), Path(file_path).name,
                )
                return result

        # Layer 4: LLM-assisted (optional)
        if llm_caller is not None and content:
            result = await LlmTocExtractor.extract(content, llm_caller)
            if cls._is_sufficient(result):
                logger.info(
                    "TOC Layer 4 (LLM): %d entries for %s",
                    len(result.entries), Path(file_path).name,
                )
                return result

        logger.debug(
            "TOC extraction: no layer produced sufficient entries for %s",
            Path(file_path).name,
        )
        return None

    @classmethod
    def _is_sufficient(cls, result: Optional[TocResult]) -> bool:
        """Check whether a TocResult has enough entries to be useful."""
        if result is None:
            return False
        return len(result.entries) >= cls._MIN_ENTRIES_THRESHOLD

    # ------------------------------------------------------------------ #
    #  Character position resolution                                      #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _resolve_char_positions(
        entries: List[TOCEntry],
        content: str,
    ) -> List[TOCEntry]:
        """Resolve character start/end positions for TOC entries in content.

        Searches for each entry's title in the content text using
        case-insensitive matching, progressing forward to avoid duplicate
        matches.  Sets char_end to the start of the next entry (or
        len(content) for the last entry).

        Also recurses into children to resolve their positions.

        Args:
            entries: Flat list of TOCEntry to resolve.
            content: Full extracted text to search within.

        Returns:
            The same list with char_start and char_end populated.
        """
        if not content or not entries:
            return entries

        content_lower = content.lower()
        search_from = 0

        # Collect all entries in document order (top-level + children)
        flat: List[TOCEntry] = []
        TOCExtractor._flatten_entries(entries, flat)

        # Pass 1: resolve char_start for each entry
        for entry in flat:
            title_lower = entry.title.lower().strip()
            if not title_lower:
                entry.char_start = search_from
                continue
            # Normalise whitespace for fuzzy matching
            title_normalised = re.sub(r"\s+", " ", title_lower)
            pos = content_lower.find(title_normalised, search_from)
            if pos < 0:
                pos = content_lower.find(title_lower, search_from)
            if pos >= 0:
                entry.char_start = pos
                search_from = pos + len(title_lower)
            else:
                pos = content_lower.find(title_normalised)
                if pos < 0:
                    pos = content_lower.find(title_lower)
                if pos >= 0:
                    entry.char_start = pos
                else:
                    entry.char_start = search_from

        # Pass 2: resolve char_end as start of next entry (or len(content))
        for i in range(len(flat) - 1):
            flat[i].char_end = flat[i + 1].char_start
        if flat:
            flat[-1].char_end = len(content)

        return entries

    @staticmethod
    def _flatten_entries(
        entries: List[TOCEntry],
        flat: List[TOCEntry],
    ) -> None:
        """Flatten nested TOCEntry tree into document-order list."""
        for entry in entries:
            flat.append(entry)
            if entry.children:
                TOCExtractor._flatten_entries(entry.children, flat)

    @staticmethod
    def _count_entries(entries: List[TOCEntry]) -> int:
        """Count total entries including nested children."""
        count = 0
        for entry in entries:
            count += 1
            if entry.children:
                count += TOCExtractor._count_entries(entry.children)
        return count
